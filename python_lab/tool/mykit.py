#!/usr/bin/python3
# -*- coding: utf-8 -*-

'''
    module description
    date: 2020/7/10
'''

from docx import Document
import requests
import time
from bs4 import BeautifulSoup
from urllib import request
import re

#读写word文档
def read_file(filename):
    doc = Document(filename)
    # 每一段的内容
    #for para in doc.paragraphs:
    #    print(para.text)

    # 每一段的编号、内容
    #for i in range(len(doc.paragraphs)):
    #    print(str(i), doc.paragraphs[i].text)

    '''for i, p in enumerate(doc.paragraphs):
        print(str(i) + ": " + str(p.text))
        p.text = (p.text).split(' ')[0] + " -DEF- "  #整行替换，会截掉后面的内容
        doc.save(filename)'''
    #取每行的第一个空格前的单词，拼接一个字符
    for p in doc.paragraphs:
        line_list = (p.text).split(' ')
        word = line_list[0]
        sign = get_oxford_phon(word)
        line_list[0] = word + ' ' + sign
        print(line_list[0])
        p.text = ' '.join(line_list)
        time.sleep(0.1)
    doc.save(filename)

#给word文档中的单词加音标，每行第一个空格前的部分为单词
def add_phon(filename):
    doc = Document(filename)
    for p in doc.paragraphs:
        line_list = (p.text).split(' ')
        word = line_list[0]
        sign = get_oxford_phon(word)
        line_list[0] = word + ' ' + sign
        print(line_list[0])
        p.text = ' '.join(line_list)
        time.sleep(0.1)
    doc.save(filename)

#抓取oxford音标
def get_oxford_phon(word):
    url = "https://www.oxfordlearnersdictionaries.com/definition/english/" + word
    wbdata = requests.get(url,headers={'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_12_1) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/55.0.2883.75 Safari/537.36'}).text
    soup = BeautifulSoup(wbdata,'html.parser')
    print(word)
    phon_html = soup.select("span.phon")
    print(phon_html)
    result = ''
    if phon_html is None: return result
    for n in phon_html:
        phon = n.get_text()
        result = result + phon
        if result != "": break
    print(result)
    return result


#抓取有道音标
def get_youdao_phon(word):
    url = 'http://dict.youdao.com/search?q=' + word + '&keyfrom=dict.index'
    response = request.urlopen(url)
    html = response.read()
    soup = BeautifulSoup(html)
    #print(soup)
    phon_html = soup.select("span.phonetic")
    print(phon_html)
    result = ''
    if phon_html is None: return result
    for n in phon_html:
        phon = n.get_text()
        result = result + phon
        if result != "": break
    print(result)
    example_html = soup.select("div.examples")  #div.examples > p
    print(example_html)
    return result

#抓取百度音标
def get_baidu_phon(word):
    url = 'https://fanyi.baidu.com/#en/zh/' + word
    response = request.urlopen(url)
    html = response.read()
    soup = BeautifulSoup(html)
    print(soup)  #<span class="phonetic-transcription"> <span>英</span> <b>[ˈriːsnt]</b>  <a href="javascript:void(0);" data-sound-lan="uk&amp;lock" data-sound-text="recent" class="op-sound"> <span class="icon-sound sound-btn"></span> </a> <a href="javascript:void(0);" data-sound-lan="uk&amp;lock" data-sound-text="recent" data-hover-tip-text="复读" class="op-repeat data-hover-tip"> <span class="icon-repeat sound-btn"></span> </a>  </span>
    phon_html = soup.select("span.phonetic-transcription > b")
    print(phon_html)
    result = ''
    if phon_html is None: return result
    for n in phon_html:
        phon = n.get_text()
        result = result + phon
        if result != "": break
    print(result)
    return result

if __name__ == '__main__':
    filename = "D:/Temp/测试文档.docx"
    #read_file(filename)
    #get_oxford_phon("recent")
    get_baidu_phon("attack")
    #get_youdao_phon("attack")