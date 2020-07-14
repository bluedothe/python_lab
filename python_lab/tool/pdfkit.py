#!/usr/bin/python3
# -*- coding: utf-8 -*-

'''
    module description
    date: 2020/7/13
'''

import os
from configparser import ConfigParser
from io import StringIO
from io import open
from concurrent.futures import ProcessPoolExecutor

from pdfminer.pdfinterp import process_pdf
from pdfminer.converter import TextConverter
#from pdfminer.layout import LAParams
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from pdfminer.layout import *

from docx import Document

from config import config


def read_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        resource_manager = PDFResourceManager()
        return_str = StringIO()
        lap_params = LAParams()

        device = TextConverter(
            resource_manager, return_str, laparams=lap_params)
        process_pdf(resource_manager, device, file)
        device.close()

        content = return_str.getvalue()
        return_str.close()
        return content


def save_text_to_word(content, file_path):
    doc = Document()
    for line in content.split('\n'):
        paragraph = doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)

#这个函数是需要自己实现的，目的是移除控制字符（换行符、制表符、转义符等），因为python-docx是不支持控制字符写入的,
# 控制字符就是ASCII码在32以下的，所以我们使用str的translate方法，把32以下的字符移除就可以。
def remove_control_characters(content):
    mpa = dict.fromkeys(range(32))
    return content.translate(mpa)

#单个文件转换
def pdf_to_word(pdf_file_path, word_file_path):
    content = read_from_pdf(pdf_file_path)
    save_text_to_word(content, word_file_path)

#多线程对多个文件进行转换，需要在config文件中设置转换的目录
def batch_transform():
    config_parser = ConfigParser()
    config_parser.read('config.cfg')
    config = config_parser['default']

    tasks = []
    with ProcessPoolExecutor(max_workers=int(config['max_worker'])) as executor:
        for file in os.listdir(config['pdf_folder']):
            extension_name = os.path.splitext(file)[1]
            if extension_name != '.pdf':
                continue
            file_name = os.path.splitext(file)[0]
            pdf_file = config['pdf_folder'] + '/' + file
            word_file = config['word_folder'] + '/' + file_name + '.docx'
            print('正在处理: ', file)
            result = executor.submit(pdf_to_word, pdf_file, word_file)
            tasks.append(result)
    while True:
        exit_flag = True
        for task in tasks:
            if not task.done():
                exit_flag = False
        if exit_flag:
            print('完成')
            exit(0)

#pdf转txt，不处理图片
def parse(pdffilename,txtfilename):
    # rb以二进制读模式打开本地pdf文件
    fn = open(pdffilename, 'rb')
    # 创建一个pdf文档分析器
    parser = PDFParser(fn)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 连接分析器 与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)

    # 提供初始化密码doc.initialize("lianxipython")
    # 如果没有密码 就创建一个空的字符串
    doc.initialize("")
    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        print("============该文件不能转换==================")
        raise PDFTextExtractionNotAllowed

    else:
        # 创建PDf资源管理器
        resource = PDFResourceManager()
        # 创建一个PDF参数分析器
        laparams = LAParams()
        # 创建聚合器,用于读取文档的对象
        device = PDFPageAggregator(resource, laparams=laparams)
        # 创建解释器，对文档编码，解释成Python能够识别的格式
        interpreter = PDFPageInterpreter(resource, device)
        # 循环遍历列表，每次处理一页的内容
        # doc.get_pages() 获取page列表
        for page in doc.get_pages():
            # 利用解释器的process_page()方法解析读取单独页数
            interpreter.process_page(page)
            # 使用聚合器get_result()方法获取内容
            layout = device.get_result()
            # 这里layout是一个LTPage对象,里面存放着这个page解析出的各种对象
            for out in layout:
                # 判断是否含有get_text()方法，获取我们想要的文字
                if hasattr(out, "get_text"):
                    print(out.get_text())
                    with open(txtfilename, 'a', encoding='utf-8') as f:
                        f.write(out.get_text() + '\n')


# 解析pdf文件，处理图片，不能处理扫描件生成的pdf文档
def parse_pic(pdffilename,docfilename):
    fp = open(pdffilename, 'rb')  # 以二进制读模式打开
    # 用文件对象来创建一个pdf文档分析器
    parser = PDFParser(fp)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 连接分析器 与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)

    # 提供初始化密码
    # 如果没有密码 就创建一个空的字符串
    doc.initialize()

    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        # 创建PDf 资源管理器 来管理共享资源
        rsrcmgr = PDFResourceManager()
        # 创建一个PDF设备对象
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解释器对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        # 用来计数页面，图片，曲线，figure，水平文本框等对象的数量
        num_page, num_image, num_curve, num_figure, num_TextBoxHorizontal = 0, 0, 0, 0, 0

        # 循环遍历列表，每次处理一个page的内容
        for page in doc.get_pages():  # doc.get_pages() 获取page列表
            num_page += 1  # 页面增一
            interpreter.process_page(page)
            # 接受该页面的LTPage对象
            layout = device.get_result()
            for x in layout:
                if isinstance(x, LTImage):  # 图片对象
                    num_image += 1
                if isinstance(x, LTCurve):  # 曲线对象
                    num_curve += 1
                if isinstance(x, LTFigure):  # figure对象
                    num_figure += 1
                if isinstance(x, LTTextBoxHorizontal):  # 获取文本内容
                    num_TextBoxHorizontal += 1  # 水平文本框对象增一
                    # 保存文本内容
                    with open(docfilename, 'a', encoding='utf-8') as f:  # 生成doc文件的文件名及路径
                        results = x.get_text()
                        f.write(results)
                        f.write('\n')
        print('对象数量：\n', '页面数：%s\n' % num_page, '图片数：%s\n' % num_image, '曲线数：%s\n' % num_curve, '水平文本框：%s\n'
              % num_TextBoxHorizontal)

if __name__ == '__main__':
    pdf_file_path = 'G:/english/word/短单词/六个字母的英文单词.pdf'
    word_file_path = 'G:/english/word/短单词/六个字母的英文单词带音标.docx'
    txt_file_path = 'G:/english/word/短单词/六个字母的英文单词带音标.txt'
    #pdf_to_word(pdf_file_path, word_file_path)
    #parse(config.sample_file_path + '/WSJD2017000308.PDF',config.sample_file_path + '/WSJD2017000308.txt')
    parse_pic(config.sample_file_path + '/春季高考面向中职毕业生志愿网上填报系统使用说明（2020-6-17版）.PDF', config.sample_file_path + '/春季高考面向中职毕业生志愿网上填报系统使用说明（2020-6-17版）.docx')